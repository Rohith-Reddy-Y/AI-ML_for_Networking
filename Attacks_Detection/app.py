from urllib.parse import unquote, unquote_plus
from flask import Flask, request, render_template, jsonify
import pandas as pd
import numpy as np
from sklearn.linear_model import LogisticRegression
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import StandardScaler
import pickle
from gensim.models.doc2vec import Doc2Vec
from nltk.tokenize import word_tokenize
import nltk
import os
import warnings
import requests
from dotenv import load_dotenv
import datetime
import io
import csv
import PyPDF2
from docx import Document

warnings.filterwarnings("ignore", category=UserWarning)

# Try to import optional dependencies
try:
    from flask_limiter import Limiter
    from flask_limiter.util import get_remote_address
    from flask_caching import Cache
    LIMITER_AVAILABLE = True
except ImportError:
    LIMITER_AVAILABLE = False
    print("Warning: flask_limiter or flask_caching not available - running without rate limiting")

try:
    import google.generativeai as genai
    from google.generativeai.types import GenerationConfig
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False
    print("Warning: google-generativeai package not available - Gemini features disabled")

# Download NLTK tokenizer data
try:
    nltk.download('punkt')
except Exception as e:
    print(f"Error downloading NLTK data: {e}")
    exit(1)

# Load environment variables
load_dotenv()

app = Flask(__name__)

# Configure rate limiting if available
if LIMITER_AVAILABLE:
    limiter = Limiter(
        app=app,
        key_func=get_remote_address,
        default_limits=["200 per day", "50 per hour"]
    )
    cache = Cache(config={'CACHE_TYPE': 'SimpleCache'})
    cache.init_app(app)
else:
    # Create dummy limiter decorator
    def limiter(*args, **kwargs):
        def decorator(f):
            return f
        return decorator

GEMINI_API_KEY = os.getenv("GOOGLE_API_KEY") or None
HF_API_TOKEN = os.getenv("HF_API_TOKEN") or None

# Configure Gemini SDK if available
if GEMINI_AVAILABLE and GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)
    gemini_model = genai.GenerativeModel("gemini-1.5-flash")
else:
    gemini_model = None
    print("Warning: Gemini AI not available - using fallback threat analysis")

# SQL/XSS feature detection setup
badwords = ['sleep', 'drop', 'uid', 'uname', 'select', 'waitfor', 'delay', 
            'system', 'union', 'order by', 'group by', 'insert', 'update', 'delete']
sql_keywords = ['or', 'and', 'union', 'select', 'insert', 'update', 'delete']

# Initialize detection log
DETECTION_LOG = "detection_log.txt"
if not os.path.exists(DETECTION_LOG):
    with open(DETECTION_LOG, 'w', encoding='utf-8') as f:
        f.write("")

def log_detection(method, query, sql_features, sql_prediction, xss_features, xss_prediction, sql_result, xss_result, error=None):
    """Log detailed detection results to file with UI-matched formatting"""
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    log_entry = f"🛡️ SQLi & XSS Threat Detector\n"
    log_entry += f"Request Method: {method}\n"
    log_entry += f"Input Query: {query}\n"
    log_entry += f"SQL Injection: {sql_result}\n"
    log_entry += f"XSS: {xss_result}\n\n"
    
    log_entry += f"Timestamp: {timestamp}\n"
    log_entry += f"SQL Features: {sql_features}\n"
    log_entry += f"SQL Prediction: {sql_prediction}\n"
    log_entry += f"XSS Features: {xss_features}\n"
    log_entry += f"XSS Prediction: {xss_prediction}\n"
    
    if error:
        log_entry += f"Error: {error}\n"
    
    log_entry += "\n" + "="*80 + "\n\n"
    
    with open(DETECTION_LOG, 'a', encoding='utf-8') as f:
        f.write(log_entry)

def process_csv(file):
    """Process CSV file and extract queries"""
    try:
        content = file.read().decode('utf-8')
        csv_reader = csv.reader(io.StringIO(content))
        next(csv_reader)  # Skip header
        queries = [row[0] for row in csv_reader if row and row[0].strip()]  # Filter out empty rows
        return queries
    except Exception as e:
        print(f"Error processing CSV: {e}")
        return []

def process_pdf(file):
    """Process PDF file and extract text"""
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() or ""
        # Extract potential queries (simple heuristic)
        queries = []
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        for line in lines:
            if any(kw.lower() in line.lower() for kw in sql_keywords + ['script', '<script']):
                queries.append(line)
        return queries if queries else lines  # Return all lines if no obvious queries found
    except Exception as e:
        print(f"Error processing PDF: {e}")
        return []

def process_docx(file):
    """Process Word document and extract text"""
    try:
        doc = Document(file)
        queries = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                if any(kw.lower() in text.lower() for kw in sql_keywords + ['script', '<script']):
                    queries.append(text)
        return queries if queries else [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    except Exception as e:
        print(f"Error processing DOCX: {e}")
        return []

def process_text_file(file):
    """Process plain text file"""
    try:
        content = file.read().decode('utf-8')
        return [line.strip() for line in content.split('\n') if line.strip()]
    except Exception as e:
        print(f"Error processing text file: {e}")
        return []

def extract_sql_features(input_data):
    if not input_data or not isinstance(input_data, str):
        return [0]*6
    try:
        s = unquote_plus(input_data.lower())
        single_q = sum(1 for i, c in enumerate(s) if c == "'" and any(
            kw in s[max(0, i-10):i+10] for kw in sql_keywords))
        double_q = sum(1 for i, c in enumerate(s) if c == '"' and any(
            kw in s[max(0, i-10):i+10] for kw in sql_keywords))
        dashes = s.count("--") if any(kw in s for kw in sql_keywords) else 0
        braces = sum(1 for i, c in enumerate(s) if c == '(' and any(
            kw in s[max(0, i-10):i+10] for kw in sql_keywords))
        spaces = s.count(' ') if any(kw in s for kw in sql_keywords) else 0
        bad_count = sum(s.count(w) for w in badwords if w in s)
        return [single_q, double_q, dashes, braces, spaces, bad_count]
    except:
        return [0]*6

# Load Doc2Vec model for XSS features
try:
    d2v_model = Doc2Vec.load("Models/d2v.model")
except Exception as e:
    print("Doc2Vec model load error:", e)
    exit(1)

def getVec(text_list):
    mats = []
    for line in text_list:
        tokens = word_tokenize(unquote(line).lower()) if line else ["default"]
        try:
            vec = d2v_model.infer_vector(tokens)
        except:
            vec = [0]*20
        lower = unquote(line).lower()
        f1 = sum(lower.count(tag) for tag in ['script', '<script', 'iframe', 'onerror', 'onload'])
        f2 = sum(lower.count(m) for m in ['alert', 'eval', 'exec', 'write', 'unescape'])
        f3 = lower.count('.js')
        f4 = lower.count('javascript')
        f5 = len(lower) if (f1 or f2) else 0
        f6 = sum(lower.count(c) for c in ['<', '>', '&'] if 'script' in lower or 'javascript' in lower)
        mats.append(np.append(vec, [f1, f2, f3, f4, f5, f6]))
    return mats

def load_or_train_sql_model():
    try:
        with open("Models/SQLi_Model.pkl", "rb") as f:
            model = pickle.load(f)
            scaler = pickle.load(f)
            return model, scaler
    except:
        df = pd.read_csv("Data/Good_and_Bad_requests.csv")
        df['class'] = df['class'].str.lower().eq('bad').astype(int)
        X = df[['single_q', 'double_q', 'dashes', 'braces', 'spaces', 'badwords']].values
        y = df['class'].values
        X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)
        scaler = StandardScaler().fit(X_train)
        model = LogisticRegression().fit(scaler.transform(X_train), y_train)
        with open("Models/SQLi_Model.pkl", "wb") as f:
            pickle.dump(model, f)
            pickle.dump(scaler, f)
        return model, scaler

def load_xss_model():
    try:
        with open("Models/RandomForestClassifier.sav", "rb") as f:
            return pickle.load(f)
    except:
        print("XSS model not found or invalid.")
        exit(1)

# Load models
sql_model, sql_scaler = load_or_train_sql_model()
xss_model = load_xss_model()

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify(error="No file uploaded"), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify(error="No selected file"), 400
    
    # Determine file type and process accordingly
    filename = file.filename.lower()
    if filename.endswith('.csv'):
        queries = process_csv(file)
    elif filename.endswith('.pdf'):
        queries = process_pdf(file)
    elif filename.endswith('.docx'):
        queries = process_docx(file)
    elif filename.endswith('.txt'):
        queries = process_text_file(file)
    else:
        return jsonify(error="Unsupported file type"), 400
    
    if not queries:
        return jsonify(error="No queries found in file"), 400
    
    # Process all queries
    results = []
    sql_detected = 0
    xss_detected = 0
    start_time = datetime.datetime.now()

    for query in queries:
        try:
            dec = unquote_plus(query)
            sf = extract_sql_features(dec)
            xf = getVec([dec])[0]
            sqlp = sql_model.predict(sql_scaler.transform([sf]))[0] if sum(sf) > 0 else 0
            xssp = xss_model.predict([xf])[0]
            
            sql_result = "SQL Injection Detected" if sqlp else "No SQL Injection"
            xss_result = "XSS Detected" if xssp else "No XSS"
            if sqlp: sql_detected += 1
            if xssp: xss_detected += 1
            
            results.append({
                'query': query,
                'sql_injection': sql_result,
                'xss': xss_result
            })
            
            # Log the detection
            log_detection(
                method="FILE_UPLOAD",
                query=query,
                sql_features=sf,
                sql_prediction=sqlp,
                xss_features=xf,
                xss_prediction=xssp,
                sql_result=sql_result,
                xss_result=xss_result
            )
            
        except Exception as e:
            results.append({
                'query': query,
                'error': str(e)
            })
            log_detection(
                method="FILE_UPLOAD",
                query=query,
                sql_features=None,
                sql_prediction=None,
                xss_features=None,
                xss_prediction=None,
                sql_result="Error",
                xss_result="Error",
                error=str(e)
            )
            return jsonify({
                'progress': 100,
                'message': 'Error occurred',
                'results': results,
                'stats': {
                    'total_queries': len(queries),
                    'processing_time': str(datetime.datetime.now() - start_time),
                    'sql_detected': sql_detected,
                    'xss_detected': xss_detected
                }
            })

    # Final response
    processing_time = datetime.datetime.now() - start_time
    return jsonify({
        'progress': 100,
        'message': 'Analysis complete',
        'results': results,
        'stats': {
            'total_queries': len(queries),
            'processing_time': str(processing_time),
            'sql_detected': sql_detected,
            'xss_detected': xss_detected
        }
    })

@app.route('/detect', methods=['GET', 'POST'])
@limiter.limit("10 per minute") if LIMITER_AVAILABLE else lambda f: f
def detect():
    try:
        method = request.method
        raw = request.form.get('query') if method == 'POST' else request.args.get('query', '')
        if not raw:
            return jsonify(error="No query provided"), 400
            
        dec = unquote_plus(raw)
        sf = extract_sql_features(dec)
        xf = getVec([dec])[0]
        sqlp = sql_model.predict(sql_scaler.transform([sf]))[0] if sum(sf) > 0 else 0
        xssp = xss_model.predict([xf])[0]
        
        sql_result = "SQL Injection Detected" if sqlp else "No SQL Injection"
        xss_result = "XSS Detected" if xssp else "No XSS"
        
        log_detection(
            method=method,
            query=raw,
            sql_features=sf,
            sql_prediction=sqlp,
            xss_features=xf,
            xss_prediction=xssp,
            sql_result=sql_result,
            xss_result=xss_result
        )
        
        return jsonify(
            sql_injection=sql_result,
            xss=xss_result
        )
        
    except Exception as e:
        error_msg = str(e)
        log_detection(
            method=request.method,
            query=request.form.get('query') if request.method == 'POST' else request.args.get('query', ''),
            sql_features=None,
            sql_prediction=None,
            xss_features=None,
            xss_prediction=None,
            sql_result="Error",
            xss_result="Error",
            error=error_msg
        )
        return jsonify(error=error_msg), 500

@app.route('/ai-threats', methods=['POST'])
@limiter.limit("5 per minute") if LIMITER_AVAILABLE else lambda f: f
def ai_threats():
    query = request.json.get('query', '').strip()
    if not query:
        return jsonify(error="Missing query input"), 400

    # Try Gemini first if available
    if GEMINI_AVAILABLE and GEMINI_API_KEY:
        try:
            prompt = f"""
Analyze this potentially malicious input: `{query}`

If it resembles SQLi or XSS, list 3-5 specific real-world threats it could cause. 
Format your response with:
- [Emoji] Threat description 1
- [Emoji] Threat description 2
- [Emoji] Threat description 3

Focus on technical impacts and be concise."""
            
            resp = gemini_model.generate_content(
                contents=prompt,
                generation_config=GenerationConfig(
                    temperature=0.3,
                    top_p=0.9,
                    max_output_tokens=300
                )
            )
            
            if hasattr(resp, 'text'):
                text = resp.text
            elif hasattr(resp, 'candidates') and resp.candidates:
                text = resp.candidates[0].content.parts[0].text
            else:
                text = ""
            
            if text:
                return jsonify(threats=text.strip())
        except Exception as ge:
            print("❌ Gemini error:", ge)

    # Fallback to Hugging Face
    if HF_API_TOKEN:
        try:
            prompt = f"Analyze security threats for: {query}. List 3-5 specific risks with emojis."
            hf_res = requests.post(
                "https://api-inference.huggingface.co/models/google/gemma-7b",
                headers={"Authorization": f"Bearer {HF_API_TOKEN}"},
                json={
                    "inputs": prompt,
                    "parameters": {
                        "max_length": 300,
                        "temperature": 0.4,
                        "return_full_text": False
                    }
                },
                timeout=15
            )
            data = hf_res.json()
            text = data[0]['generated_text'] if isinstance(data, list) else ""
            if text:
                return jsonify(threats=text.strip())
        except Exception as he:
            print("⚠️ HF fallback error:", he)

    # Final fallback
    default_threats = """
- 🔓 Unauthorized database access
- 🗑️ Data deletion or corruption
- 📜 Database schema exposure
- 📊 Sensitive data exfiltration
- ⚠️ Complete system compromise
"""
    return jsonify(threats=default_threats.strip())

@app.route('/favicon.ico')
def favicon():
    return '', 204

@app.route('/explain-threat', methods=['POST'])
def explain_threat():
    data = request.json
    question = data.get('question', '').strip()
    current_threats = data.get('current_threats', '')
    
    if not question:
        return jsonify(error="No question provided"), 400
    
    try:
        # Try Gemini first if available
        if GEMINI_AVAILABLE and GEMINI_API_KEY:
            prompt = f"""
You are a cybersecurity expert explaining SQL injection threats to a user.
The user has seen these potential threats:
{current_threats}

They are asking: {question}

Provide a detailed but concise explanation (2-3 paragraphs max) in simple terms.
Include examples if helpful.
"""
            resp = gemini_model.generate_content(
                contents=prompt,
                generation_config=GenerationConfig(
                    temperature=0.3,
                    top_p=0.9,
                    max_output_tokens=500
                )
            )
            
            if hasattr(resp, 'text'):
                text = resp.text
            elif hasattr(resp, 'candidates') and resp.candidates:
                text = resp.candidates[0].content.parts[0].text
            else:
                text = ""
            
            if text:
                return jsonify(explanation=text.strip())
        
        # Fallback to simple explanations
        explanations = {
            "data breach": "A data breach occurs when attackers access sensitive information they shouldn't have access to. In SQL injection, this happens when malicious queries extract data from tables containing user credentials, personal information, or other private data.",
            "command execution": "Some SQL databases allow running system commands through special functions. Attackers can chain these to your vulnerable query to execute any command on your server, potentially taking full control.",
            "file overwrite": "Certain SQL functions let you write files. Attackers can overwrite critical files like ASP scripts to insert malicious code that gives them persistent access to your system.",
            "denial of service": "DoS attacks make your database unresponsive. Attackers can craft queries that consume all resources - like endless joins or recursive queries - preventing legitimate users from accessing data."
        }
        
        for term in explanations:
            if term in question.lower():
                return jsonify(explanation=explanations[term])
        
        return jsonify(explanation="This threat involves potential security risks from the SQL injection. For more details, consult a cybersecurity professional.")
        
    except Exception as e:
        return jsonify(error=str(e)), 500

if __name__ == "__main__":
    app.run(debug=True)